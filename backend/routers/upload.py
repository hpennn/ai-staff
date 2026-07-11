"""
文件上传路由 - 解析文件内容为文本行（用于知识库/关键词导入）
支持: txt, csv, md, pdf, docx, xlsx
"""
from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel
from typing import Optional, List
import os
import tempfile
import logging
import httpx

router = APIRouter()

logger = logging.getLogger(__name__)

# 允许的文件类型和大小限制（10MB）
ALLOWED_EXTENSIONS = {'.txt', '.csv', '.md', '.pdf', '.docx', '.xlsx', '.json', '.log'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


class FileContentResponse(BaseModel):
    lines: List[str]
    total_lines: int
    file_name: str
    file_type: str


def read_text_file(filepath: str, encoding: str = 'utf-8') -> List[str]:
    """读取纯文本文件，返回非空行列表"""
    try:
        with open(filepath, 'r', encoding=encoding, errors='ignore') as f:
            content = f.read()
        # 按行分割，过滤空白行
        lines = [line.strip() for line in content.splitlines() if line.strip()]
        return lines
    except Exception as e:
        logger.error(f"读取文本文件失败: {e}")
        raise HTTPException(status_code=400, detail=f"文件读取失败: {str(e)}")


def read_csv_file(filepath: str) -> List[str]:
    """读取CSV文件，返回所有非空单元格内容"""
    try:
        import csv
        lines = []
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                for cell in row:
                    cell = cell.strip()
                    if cell:
                        lines.append(cell)
        return lines
    except Exception as e:
        logger.error(f"读取CSV文件失败: {e}")
        raise HTTPException(status_code=400, detail=f"CSV解析失败: {str(e)}")


def read_pdf_file(filepath: str) -> List[str]:
    """读取PDF文件，返回文本行"""
    try:
        import PyPDF2
        text_lines = []
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    lines = [line.strip() for line in text.splitlines() if line.strip()]
                    text_lines.extend(lines)
        return text_lines
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF解析库未安装，请联系管理员")
    except Exception as e:
        logger.error(f"读取PDF文件失败: {e}")
        raise HTTPException(status_code=400, detail=f"PDF解析失败: {str(e)}")


def read_docx_file(filepath: str) -> List[str]:
    """读取Word文档，返回文本行"""
    try:
        from docx import Document
        doc = Document(filepath)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # 也读取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
        return lines
    except ImportError:
        raise HTTPException(status_code=500, detail="Word文档解析库未安装，请联系管理员")
    except Exception as e:
        logger.error(f"读取Word文件失败: {e}")
        raise HTTPException(status_code=400, detail=f"Word文档解析失败: {str(e)}")


def read_xlsx_file(filepath: str) -> List[str]:
    """读取Excel文件，返回所有非空单元格内容"""
    try:
        import openpyxl
        lines = []
        wb = openpyxl.load_workbook(filepath, read_only=True)
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        val = str(cell.value).strip()
                        if val:
                            lines.append(val)
        wb.close()
        return lines
    except ImportError:
        raise HTTPException(status_code=500, detail="Excel解析库未安装，请联系管理员")
    except Exception as e:
        logger.error(f"读取Excel文件失败: {e}")
        raise HTTPException(status_code=400, detail=f"Excel解析失败: {str(e)}")


def read_json_file(filepath: str) -> List[str]:
    """读取JSON文件，尝试提取文本内容"""
    try:
        import json
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        lines = []
        if isinstance(data, list):
            for item in data:
                if isinstance(item, str):
                    if item.strip():
                        lines.append(item.strip())
                elif isinstance(item, dict):
                    # 尝试提取常见文本字段
                    for key in ['text', 'content', 'answer', 'question', 'q', 'a', 'name', 'description']:
                        if key in item and isinstance(item[key], str):
                            val = item[key].strip()
                            if val:
                                lines.append(val)
        elif isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, str) and value.strip():
                    lines.append(f"{key}: {value.strip()}")
        return lines
    except Exception as e:
        logger.error(f"读取JSON文件失败: {e}")
        raise HTTPException(status_code=400, detail=f"JSON解析失败: {str(e)}")




class ImportUrlRequest(BaseModel):
    url: str


class ImportUrlResponse(BaseModel):
    lines: List[str]
    total_lines: int
    url: str


@router.post("/import-url", response_model=ImportUrlResponse)
@router.post("/import-url", response_model=ImportUrlResponse)
async def import_url(request: ImportUrlRequest):
    """从网页URL导入文本内容到知识库"""
    import re as regex_module
    from html import unescape as html_unescape

    url = request.url.strip()
    if not url:
        raise HTTPException(status_code=400, detail="请输入网址")

    if not url.startswith(('http://', 'https://')):
        raise HTTPException(status_code=400, detail="请输入有效的网址（以http://或https://开头）")

    try:
        async with httpx.AsyncClient(timeout=15.0, follow_redirects=True) as client:
            resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0 (compatible; AI-Staff-Bot/1.0)"})
            resp.raise_for_status()
            html_content = resp.text
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="网页加载超时，请稍后重试")
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=400, detail=f"网页访问失败: HTTP {e.response.status_code}")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"无法访问该网址: {str(e)}")

    # Remove script and style tags with content
    html_content = regex_module.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=regex_module.DOTALL | regex_module.IGNORECASE)
    html_content = regex_module.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=regex_module.DOTALL | regex_module.IGNORECASE)
    html_content = regex_module.sub(r'<noscript[^>]*>.*?</noscript>', '', html_content, flags=regex_module.DOTALL | regex_module.IGNORECASE)

    # Remove HTML tags
    text = regex_module.sub(r'<[^>]+>', ' ', html_content)

    # Decode HTML entities
    text = html_unescape(text)

    # Clean up whitespace
    text = regex_module.sub(r'[ \t]+', ' ', text)
    text = regex_module.sub(r'\n{2,}', '\n', text)

    # Split into lines and filter
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    # Limit to 200 lines
    if len(lines) > 200:
        lines = lines[:200]

    return ImportUrlResponse(
        lines=lines,
        total_lines=len(lines),
        url=url
    )


@router.post("/upload/file", response_model=FileContentResponse)
async def upload_file(file: UploadFile = File(...)):
    """
    上传文件并解析为文本行
    支持: txt, csv, md, pdf, docx, xlsx, json, log
    """
    # 检查文件扩展名
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400, 
            detail=f"不支持的文件格式: {file_ext}，支持: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    # 检查文件大小
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400, 
            detail=f"文件过大，最大支持 {MAX_FILE_SIZE // 1024 // 1024}MB"
        )

    # 保存到临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        # 根据文件类型读取内容
        if file_ext in ['.txt', '.md', '.log']:
            lines = read_text_file(tmp_path)
        elif file_ext == '.csv':
            lines = read_csv_file(tmp_path)
        elif file_ext == '.pdf':
            lines = read_pdf_file(tmp_path)
        elif file_ext == '.docx':
            lines = read_docx_file(tmp_path)
        elif file_ext == '.xlsx':
            lines = read_xlsx_file(tmp_path)
        elif file_ext == '.json':
            lines = read_json_file(tmp_path)
        else:
            lines = read_text_file(tmp_path)

        # 限制返回行数（最多500行）
        if len(lines) > 500:
            lines = lines[:500]

        return FileContentResponse(
            lines=lines,
            total_lines=len(lines),
            file_name=file.filename,
            file_type=file_ext[1:]  # 去掉点
        )
    finally:
        # 清理临时文件
        try:
            os.unlink(tmp_path)
        except:
            pass
