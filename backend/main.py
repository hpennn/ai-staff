import sys
import os

# Add backend directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from fastapi import FastAPI
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from database import init_db
from routers import chat, staff, webhook, auth, admin, subscription, upload, broadcast, schedule, skills, knowledge, workflow, computer_ctrl, agent_chat

# Initialize database
init_db()

# Initialize credits database
from credits_database import init_credits_db
init_credits_db()

# Initialize auth database (user phone login)
from auth_database import init_auth_db
init_auth_db()

# Load skills engine
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from skills.registry import registry
registry.load_preset_skills()

app = FastAPI(title="智能体工作台", version="2.2.0")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# API routes - auth is public
app.include_router(auth.router, prefix="/api")
# Chat and webhook are public
app.include_router(chat.router, prefix="/api")
app.include_router(webhook.router, prefix="/api")
# Staff routes (some public, some require auth internally)
app.include_router(staff.router, prefix="/api")
# Admin routes (all require auth)
app.include_router(admin.router, prefix="/api")
app.include_router(subscription.router, prefix="/api")
app.include_router(upload.router, prefix="/api")
# Feature 5: Broadcast
app.include_router(broadcast.router, prefix="/api")
# Feature 7: Schedule
app.include_router(schedule.router, prefix="/api")
# Skills engine
app.include_router(skills.router, prefix="/api")
# Knowledge base
app.include_router(knowledge.router, prefix="/api")
# Workflow orchestration
app.include_router(workflow.router, prefix="/api")
# Computer control
app.include_router(computer_ctrl.router, prefix="/api")
# Agent chat (frontend agents)
app.include_router(agent_chat.router, prefix="/api")
# Credits system
from routers.credits import router as credits_router
app.include_router(credits_router, prefix="/api")



@app.get("/chat/{staff_id}")
async def chat_page(staff_id: int):
    """Serve chat widget page for embedding"""
    from database import get_db
    conn = get_db()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM staff WHERE id = ?", (staff_id,))
    staff = cursor.fetchone()
    conn.close()
    if not staff:
        from fastapi.responses import HTMLResponse
        return HTMLResponse("<h1>员工不存在</h1>", status_code=404)
    
    staff_dict = dict(staff)
    from fastapi.responses import HTMLResponse
    
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
<title>{staff_dict['name']}</title>
<script src="https://cdn.tailwindcss.com"></script>
<style>
* {{ -webkit-tap-highlight-color: transparent; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; margin: 0; padding: 0; }}
.chat-messages {{ flex: 1; overflow-y: auto; }}
.msg-bubble {{ max-width: 80%; word-wrap: break-word; }}
.fade-in {{ animation: fadeIn 0.3s ease; }}
@keyframes fadeIn {{ from {{ opacity: 0; transform: translateY(10px); }} to {{ opacity: 1; transform: translateY(0); }} }}
input:focus, textarea:focus {{ outline: none; box-shadow: 0 0 0 2px rgba(99,102,241,0.3); }}
</style>
</head>
<body class="bg-slate-50 h-screen overflow-hidden">
<div class="h-full flex flex-col max-w-lg mx-auto bg-white relative">
    <div class="flex items-center px-3 py-3 bg-white border-b border-slate-100 shadow-sm">
        <div class="w-8 h-8 rounded-full flex items-center justify-center text-sm font-semibold text-white mr-2" style="background:{staff_dict.get('avatar_color','#6366f1')}">
{staff_dict['name'][0]}</div>
        <div class="flex-1">
            <h2 class="font-semibold text-slate-800 text-sm">{staff_dict['name']}</h2>
            <p class="text-xs text-slate-400 truncate">{staff_dict.get('role_description','')}</p>
        </div>
    </div>
    <div id="chat-messages" class="chat-messages px-4 py-4 space-y-3"></div>
    <div id="rating-area" class="hidden bg-white border-t border-slate-100 px-4 py-2 flex items-center justify-center gap-4">
        <span class="text-xs text-slate-400">对本次对话评价：</span>
        <button onclick="rateConversation('good')" id="btn-good" class="px-3 py-1 rounded-lg border border-slate-200 hover:bg-green-50 text-sm transition-colors">👍 好评</button>
        <button onclick="rateConversation('bad')" id="btn-bad" class="px-3 py-1 rounded-lg border border-slate-200 hover:bg-red-50 text-sm transition-colors">👎 差评</button>
    </div>
    <div class="bg-white border-t border-slate-100 px-3 py-2 flex items-end gap-2">
        <textarea id="chat-input" rows="1" placeholder="输入消息..." class="flex-1 resize-none border border-slate-200 rounded-xl px-3 py-2 text-sm focus:border-indigo-400 transition-colors" oninput="this.style.height='auto';this.style.height=Math.min(this.scrollHeight,100)+'px'" onkeydown="if(event.key==='Enter'&&!event.shiftKey){{event.preventDefault();sendMessage()}}"></textarea>
        <button onclick="sendMessage()" class="w-9 h-9 bg-indigo-500 text-white rounded-full flex items-center justify-center hover:bg-indigo-600 transition-colors shadow-sm flex-shrink-0">
            <svg class="w-4 h-4" fill="none" stroke="currentColor" viewBox="0 0 24 24"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M12 19l9 2-9-18-9 18 9-2zm0 0v-8"/></svg>
        </button>
    </div>
</div>
<script>
const staffId = {staff_id};
const sessionId = 'embed-' + Date.now();
let convId = null;
let hasRated = false;
const welcome = "{staff_dict.get('welcome_message', '')}";
if (welcome) addBubble('assistant', welcome);

function addBubble(role, content) {{
    const c = document.getElementById('chat-messages');
    const d = document.createElement('div');
    d.className = 'flex ' + (role === 'user' ? 'justify-end' : 'justify-start') + ' fade-in';
    if (role === 'user') {{
        d.innerHTML = '<div class="msg-bubble bg-indigo-500 text-white rounded-2xl rounded-br-md px-4 py-2.5 text-sm shadow-sm">' + escapeHtml(content) + '</div>';
    }} else {{
        d.innerHTML = '<div class="msg-bubble bg-white text-slate-700 rounded-2xl rounded-bl-md px-4 py-2.5 text-sm shadow-sm border border-slate-100">' + escapeHtml(content) + '</div>';
    }}
    c.appendChild(d);
    setTimeout(() => c.scrollTop = c.scrollHeight, 50);
}}

function escapeHtml(t) {{ const d = document.createElement('div'); d.textContent = t; return d.innerHTML; }}

async function sendMessage() {{
    const input = document.getElementById('chat-input');
    const msg = input.value.trim();
    if (!msg) return;
    input.value = '';
    input.style.height = 'auto';
    addBubble('user', msg);
    const typing = document.createElement('div');
    typing.id = 'typing';
    typing.className = 'flex justify-start fade-in';
    typing.innerHTML = '<div class="msg-bubble bg-white text-slate-400 rounded-2xl rounded-bl-md px-4 py-2.5 text-sm shadow-sm border border-slate-100">思考中...</div>';
    document.getElementById('chat-messages').appendChild(typing);
    setTimeout(() => document.getElementById('chat-messages').scrollTop = document.getElementById('chat-messages').scrollHeight, 50);
    try {{
        const res = await fetch('/api/chat', {{
            method: 'POST',
            headers: {{'Content-Type': 'application/json'}},
            body: JSON.stringify({{staff_id: staffId, session_id: sessionId, message: msg}})
        }});
        const data = await res.json();
        typing.remove();
        addBubble('assistant', data.reply);
        // Show rating after first AI reply
        if (!hasRated) document.getElementById('rating-area').classList.remove('hidden');
    }} catch(e) {{
        typing.remove();
        addBubble('assistant', '⚠️ ' + e.message);
    }}
}}

async function rateConversation(rating) {{
    if (hasRated) return;
    try {{
        const res = await fetch('/api/rate/' + sessionId, {{
            method: 'POST',
            headers: {{'Content-Type': 'application/json'}},
            body: JSON.stringify({{rating: rating}})
        }});
        if (res.ok) {{
            hasRated = true;
            document.getElementById('btn-good').disabled = true;
            document.getElementById('btn-bad').disabled = true;
            if (rating === 'good') {{
                document.getElementById('btn-good').classList.add('bg-green-100', 'border-green-300');
                document.getElementById('btn-good').textContent = '👍 已好评';
            }} else {{
                document.getElementById('btn-bad').classList.add('bg-red-100', 'border-red-300');
                document.getElementById('btn-bad').textContent = '👎 已差评';
            }}
        }}
    }} catch(e) {{}}
}}
</script>
</body>
</html>"""
    return HTMLResponse(content=html)


@app.get("/embed/{staff_id}.js")
async def embed_script(staff_id: int):
    """Generate embed JS for floating chat bubble"""
    from fastapi.responses import Response
    js_code = f"""
(function() {{
    var staffId = {staff_id};
    var baseUrl = window.location.origin;
    
    // Create bubble button
    var bubble = document.createElement('div');
    bubble.id = 'ai-staff-bubble';
    bubble.innerHTML = '💬';
    bubble.style.cssText = 'position:fixed;bottom:24px;right:24px;width:56px;height:56px;border-radius:50%;background:#6366f1;color:white;display:flex;align-items:center;justify-content:center;font-size:24px;cursor:pointer;box-shadow:0 4px 12px rgba(99,102,241,0.4);z-index:99999;transition:transform 0.2s;';
    bubble.onmouseenter = function() {{ this.style.transform = 'scale(1.1)'; }};
    bubble.onmouseleave = function() {{ this.style.transform = 'scale(1)'; }};
    
    // Create chat window
    var chatWindow = document.createElement('div');
    chatWindow.id = 'ai-staff-chat-window';
    chatWindow.style.cssText = 'position:fixed;bottom:90px;right:24px;width:380px;height:600px;border-radius:16px;overflow:hidden;box-shadow:0 8px 32px rgba(0,0,0,0.15);z-index:99998;display:none;';
    
    var iframe = document.createElement('iframe');
    iframe.src = baseUrl + '/chat/' + staffId;
    iframe.style.cssText = 'width:100%;height:100%;border:none;';
    iframe.setAttribute('frameborder', '0');
    chatWindow.appendChild(iframe);
    
    var isOpen = false;
    bubble.onclick = function() {{
        isOpen = !isOpen;
        chatWindow.style.display = isOpen ? 'block' : 'none';
        bubble.innerHTML = isOpen ? '✕' : '💬';
    }};
    
    document.body.appendChild(bubble);
    document.body.appendChild(chatWindow);
}})();
"""
    return Response(content=js_code, media_type="application/javascript")


# 文件下载目录
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "data", "downloads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.get("/api/download/{filename}")
async def download_file(filename: str):
    from fastapi.responses import FileResponse
    filepath = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(filepath):
        return {"error": "文件不存在"}
    return FileResponse(filepath, media_type="application/octet-stream", filename=filename)

@app.post("/api/save-result")
async def save_result(data: dict):
    """保存识别结果为文件并返回下载链接，支持 txt/md/pdf/docx 四种格式"""
    import uuid
    from datetime import datetime

    fmt = data.get("format", "txt").lower()
    content = data.get("content", "")
    custom_filename = data.get("filename", "").strip()

    # 格式到扩展名映射
    ext_map = {
        "txt": "txt",
        "text": "txt",
        "md": "md",
        "markdown": "md",
        "pdf": "pdf",
        "docx": "docx",
        "word": "docx",
    }
    ext = ext_map.get(fmt, "txt")

    # 生成文件名
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    short_id = uuid.uuid4().hex[:6]
    if custom_filename:
        # 去掉用户文件名中的扩展名，统一加上
        import re
        base_name = re.sub(r'\.(txt|md|pdf|docx)$', '', custom_filename, flags=re.IGNORECASE)
        filename = f"{base_name}_{timestamp}.{ext}"
    else:
        filename = f"result_{timestamp}_{short_id}.{ext}"

    filepath = os.path.join(UPLOAD_DIR, filename)

    try:
        if ext in ("txt", "md"):
            # 纯文本 / Markdown 直接写入
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)

        elif ext == "pdf":
            # 使用 fpdf2 生成 PDF，支持中文
            from fpdf import FPDF

            class PDF(FPDF):
                def header(self):
                    pass

                def footer(self):
                    pass

            pdf = PDF()
            pdf.add_page()

            # 注册中文字体：优先使用系统 Noto CJK 字体，其次尝试常见中文字体路径
            font_paths = [
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
                "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
                "/System/Library/Fonts/PingFang.ttc",
            ]
            font_registered = False
            for fp in font_paths:
                if os.path.exists(fp):
                    try:
                        pdf.add_font("NotoSansSC", "", fp)
                        pdf.set_font("NotoSansSC", size=11)
                        font_registered = True
                        break
                    except Exception:
                        continue

            if not font_registered:
                # 兜底：使用内置字体，可能中文显示为问号，但不报错
                pdf.set_font("Helvetica", size=11)

            # 设置左右边距
            pdf.set_left_margin(15)
            pdf.set_right_margin(15)

            # 按行写入，自动换行
            for line in content.split("\n"):
                if not line:
                    pdf.ln(6)
                else:
                    # multi_cell 自动换行
                    pdf.multi_cell(0, 7, line)

            pdf.output(filepath)

        elif ext == "docx":
            # 使用 python-docx 生成 Word 文档
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn

            doc = Document()

            # 设置默认字体（中文）
            style = doc.styles["Normal"]
            font = style.font
            font.name = "微软雅黑"
            font.size = Pt(11)
            # 设置中文字体
            rpr = style.element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                rpr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 按段落写入内容
            for line in content.split("\n"):
                if not line:
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(line)

            doc.save(filepath)

    except Exception as e:
        return {"error": f"生成文件失败: {str(e)}"}, 500

    return {"filename": filename, "download_url": f"/api/download/{filename}", "format": ext}

# Static files
frontend_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "frontend")
app.mount("/static", StaticFiles(directory=frontend_path), name="static")


@app.get("/")
async def serve_index():
    return FileResponse(os.path.join(frontend_path, "index.html"))


@app.get("/manifest.json")
async def serve_manifest():
    return FileResponse(os.path.join(frontend_path, "manifest.json"))


@app.get("/sw.js")
async def serve_sw():
    return FileResponse(os.path.join(frontend_path, "sw.js"))


@app.get("/icons/{filename}")
async def serve_icon(filename: str):
    return FileResponse(os.path.join(frontend_path, "icons", filename))


@app.get("/download/apk")
async def download_apk():
    apk_path = os.path.join(frontend_path, "app.apk")
    if os.path.isfile(apk_path):
        return FileResponse(apk_path, media_type="application/vnd.android.package-archive", filename="智能工作台-v1.0.apk")
    from fastapi.responses import JSONResponse
    return JSONResponse({"error": "APK not found"}, status_code=404)




@app.get("/download/exe")
async def download_exe():
    exe_path = os.path.join(frontend_path, "app.exe")
    if os.path.isfile(exe_path):
        return FileResponse(exe_path, media_type="application/octet-stream", filename="智能工作台.exe", headers={"accept-ranges": "bytes"})
    from fastapi.responses import JSONResponse
    return JSONResponse({"error": "EXE not found"}, status_code=404)


@app.get("/download/dmg")
async def download_dmg():
    dmg_path = os.path.join(frontend_path, "app-mac.tar.gz")
    if os.path.isfile(dmg_path):
        return FileResponse(dmg_path, media_type="application/octet-stream", filename="智能工作台-mac.tar.gz", headers={"accept-ranges": "bytes"})
    from fastapi.responses import JSONResponse
    return JSONResponse({"error": "DMG not found"}, status_code=404)


@app.get("/download.html")
async def serve_download():
    return FileResponse(os.path.join(frontend_path, "download.html"))


@app.get("/health")
async def health_check():
    return {"status": "ok", "message": "智能体工作台运行中", "version": "2.2.0"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
