"""文件格式转换技能 - 图片/PDF/Word互转"""
import os
import io
import time
import base64

# 下载目录
DOWNLOAD_DIR = os.path.join(os.path.dirname(__file__), "../../static/downloads")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

SKILL_META = {
    "id": "file_converter",
    "name": "格式转换",
    "icon": "🔄",
    "description": "图片、PDF、Word文档格式互转",
    "keywords": ["转换", "格式", "pdf", "word", "图片", "docx", "互转", "转pdf", "转word"],
    "input_type": "file",
    "output_type": "file",
}

FORMAT_MAP = {
    ".jpg": "jpeg", ".jpeg": "jpeg", ".png": "png", ".bmp": "bmp",
    ".webp": "webp", ".gif": "gif", ".tiff": "tiff",
    ".pdf": "pdf", ".docx": "docx", ".doc": "docx",
}

IMAGE_FORMATS = {"jpeg", "png", "bmp", "webp", "gif", "tiff"}


def detect_format(filepath):
    """从文件扩展名识别格式"""
    ext = os.path.splitext(filepath)[1].lower()
    return FORMAT_MAP.get(ext, ext.lstrip("."))


# ========== 核心转换函数 ==========

def image_to_pdf(image_path, output_path):
    """图片 -> PDF"""
    from PIL import Image
    img = Image.open(image_path)
    if img.mode != "RGB":
        img = img.convert("RGB")
    # A4尺寸适配
    a4_w, a4_h = 2480, 3508
    ratio = min(a4_w / img.width, a4_h / img.height, 1.0)
    if ratio < 1.0:
        img = img.resize(
            (int(img.width * ratio), int(img.height * ratio)),
            Image.LANCZOS
        )
    img.save(output_path, "PDF", resolution=300)
    return {"path": output_path, "pages": 1}


def pdf_to_images(pdf_path, output_dir):
    """PDF -> 图片(每张一页PNG)"""
    import fitz  # PyMuPDF
    doc = fitz.open(pdf_path)
    pages = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=200)
        out_path = os.path.join(output_dir, f"page_{i+1}.png")
        pix.save(out_path)
        pages.append(out_path)
    doc.close()
    if pages:
        return {"path": pages[0], "paths": pages, "pages": len(pages)}
    else:
        return {"error": "PDF无内容"}


def image_to_word(image_path, output_path):
    """图片 -> Word(嵌入图片)"""
    from docx import Document
    from docx.shared import Inches
    from PIL import Image
    doc = Document()
    img = Image.open(image_path)
    w_inch = img.width / 150  # 约150dpi
    width = min(Inches(w_inch), Inches(6))
    doc.add_picture(image_path, width=width)
    doc.save(output_path)
    return {"path": output_path}


def pdf_to_word(pdf_path, output_path):
    """PDF -> Word(提取文字)"""
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Pt

    doc_pdf = fitz.open(pdf_path)
    doc_word = Document()

    style = doc_word.styles["Normal"]
    style.font.size = Pt(11)

    for page in doc_pdf:
        text = page.get_text()
        if text.strip():
            doc_word.add_paragraph(text)

    doc_pdf.close()
    doc_word.save(output_path)
    return {"path": output_path}


def word_to_pdf(docx_path, output_path):
    """Word -> PDF(提取文字重排)"""
    from docx import Document
    from fpdf import FPDF

    doc = Document(docx_path)
    pdf = FPDF()
    pdf.set_auto_page_break(True, margin_x=15, margin_y=15)
    pdf.add_page()

    # 尝试注册中文字体
    font_paths = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    ]
    font_ok = False
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdf.add_font("CJK", "", fp)
                pdf.set_font("CJK", size=11)
                font_ok = True
                break
            except Exception:
                continue

    if not font_ok:
        pdf.set_font("Helvetica", size=11)

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            pdf.ln(5)
        else:
            pdf.multi_cell(0, 7, line)

    pdf.output(output_path)
    return {"path": output_path, "pages": pdf.pages_count}


def word_to_images(docx_path, output_dir):
    """Word -> 图片(Word->PDF->图片)"""
    tmp_pdf = os.path.join(output_dir, f"_tmp_{int(time.time())}.pdf")
    result = word_to_pdf(docx_path, tmp_pdf)
    if "error" in result:
        return result
    img_result = pdf_to_images(tmp_pdf, output_dir)
    try:
        os.remove(tmp_pdf)
    except Exception:
        pass
    return img_result


def image_to_image(image_path, target_fmt, output_path):
    """图片 -> 图片(格式互转)"""
    from PIL import Image
    img = Image.open(image_path)
    save_fmt = target_fmt.upper()
    if target_fmt == "jpeg":
        img = img.convert("RGB")
    img.save(output_path, save_fmt, quality=95)
    return {"path": output_path}


# ========== 路由 ==========

def convert(src_path, target_format, output_dir):
    """根据源格式和目标格式路由到对应转换函数"""
    src_fmt = detect_format(src_path)
    tgt_fmt = target_format.lower()

    if src_fmt == tgt_fmt:
        return {"error": f"源格式和目标格式相同（均为{src_fmt}）"}

    ts = int(time.time())

    # 图片 -> PDF
    if src_fmt in IMAGE_FORMATS and tgt_fmt == "pdf":
        return image_to_pdf(
            src_path, os.path.join(output_dir, f"out_{ts}.pdf")
        )
    # 图片 -> Word
    if src_fmt in IMAGE_FORMATS and tgt_fmt == "docx":
        return image_to_word(
            src_path, os.path.join(output_dir, f"out_{ts}.docx")
        )
    # 图片 -> 图片
    if src_fmt in IMAGE_FORMATS and tgt_fmt in IMAGE_FORMATS:
        ext = ".jpg" if tgt_fmt == "jpeg" else f".{tgt_fmt}"
        return image_to_image(
            src_path, tgt_fmt, os.path.join(output_dir, f"out_{ts}{ext}")
        )
    # PDF -> 图片
    if src_fmt == "pdf" and tgt_fmt in IMAGE_FORMATS:
        return pdf_to_images(src_path, output_dir)
    # PDF -> Word
    if src_fmt == "pdf" and tgt_fmt == "docx":
        return pdf_to_word(
            src_path, os.path.join(output_dir, f"out_{ts}.docx")
        )
    # Word -> PDF
    if src_fmt == "docx" and tgt_fmt == "pdf":
        return word_to_pdf(
            src_path, os.path.join(output_dir, f"out_{ts}.pdf")
        )
    # Word -> 图片
    if src_fmt == "docx" and tgt_fmt in IMAGE_FORMATS:
        return word_to_images(src_path, output_dir)

    return {"error": f"暂不支持 {src_fmt} -> {tgt_fmt} 转换"}


def parse_target_format(text, src_fmt):
    """从用户文字解析目标格式"""
    t = text.lower()

    if any(kw in t for kw in ["转pdf", "to pdf", "生成pdf"]):
        return "pdf"
    if any(kw in t for kw in ["转word", "转docx", "to word", "生成word"]):
        return "docx"
    if any(kw in t for kw in ["转png", "to png"]):
        return "png"
    if any(kw in t for kw in ["转jpg", "转jpeg", "to jpg"]):
        return "jpeg"
    if any(kw in t for kw in ["转webp"]):
        return "webp"
    if any(kw in t for kw in ["转图片", "转成图片", "转为图片"]):
        return "png" if src_fmt == "pdf" else "jpeg"

    # 默认推断
    if src_fmt == "pdf":
        return "png"
    elif src_fmt == "docx":
        return "pdf"
    else:
        return "pdf"


# ========== 技能入口 ==========

async def execute(input_data: dict) -> dict:
    """
    输入: {"text": "转pdf", "file_path": "/tmp/xxx.png", "files": [...]}
    输出: {"file_url": "...", "content": "转换结果说明"}
    """
    text = input_data.get("text", "")
    target_format = input_data.get("target_format", "")

    # 获取源文件路径
    file_path = (
        input_data.get("file_path", "")
        or input_data.get("image_path", "")
    )
    files = input_data.get("files", [])
    if not file_path and files:
        file_path = files[0].get("filepath", "")

    if not file_path or not os.path.exists(file_path):
        return {
            "message": (
                "\U0001f504 **\u683c\u5f0f\u8f6c\u6362**\n\n"
                "\u652f\u6301\u683c\u5f0f\u4e92\u8f6c\uff1a\n\n"
                "- \U0001f4c4 **\u56fe\u7247 \u2192 PDF**\uff1aPNG/JPG/BMP\u8f6cPDF\n"
                "- \U0001f5bc\ufe0f **PDF \u2192 \u56fe\u7247**\uff1aPDF\u6bcf\u9875\u5bfc\u51faPNG\n"
                "- \U0001f4dd **\u56fe\u7247 \u2192 Word**\uff1a\u56fe\u7247\u5d4c\u5165Word\u6587\u6863\n"
                "- \U0001f4c4 **PDF \u2192 Word**\uff1a\u63d0\u53d6\u6587\u5b57\u751f\u6210Word\n"
                "- \U0001f4c4 **Word \u2192 PDF**\uff1aWord\u8f6cPDF\n"
                "- \U0001f5bc\ufe0f **Word \u2192 \u56fe\u7247**\uff1aWord\u5bfc\u51fa\u4e3a\u56fe\u7247\n\n"
                "\u8bf7\u4e0a\u4f20\u6587\u4ef6\uff0c\u5e76\u544a\u8bc9\u6211\u8981\u8f6c\u6210\u4ec0\u4e48\u683c\u5f0f\uff08\u5982\u300c\u8f6cpdf\u300d\uff09"
            )
        }

    src_fmt = detect_format(file_path)
    if not target_format:
        target_format = parse_target_format(text, src_fmt)

    tgt_fmt = target_format.lower()
    if tgt_fmt in ("jpg", "jpeg"):
        tgt_fmt = "jpeg"
    elif tgt_fmt in ("doc", "docx", "word"):
        tgt_fmt = "docx"
    elif tgt_fmt == "pdf":
        tgt_fmt = "pdf"

    result = convert(file_path, tgt_fmt, DOWNLOAD_DIR)

    if "error" in result:
        return {"error": result["error"]}

    # PDF转多张图片的特殊处理
    if tgt_fmt in IMAGE_FORMATS and "paths" in result and len(result["paths"]) > 1:
        urls = [
            f"/static/downloads/{os.path.basename(p)}"
            for p in result["paths"]
        ]
        content = (
            "\u2705 **PDF\u8f6c\u56fe\u7247\u5b8c\u6210**\n\n"
            f"\u5171 **{result['pages']}** \u9875\n\n"
            + "\n".join([
                f"[\u7b2c{i+1}\u9875]({u})"
                for i, u in enumerate(urls)
            ])
        )
        return {"file_url": urls[0], "content": content}

    # 单文件结果
    out_path = result.get("path", "")
    if not out_path or not os.path.exists(out_path):
        return {"error": "\u8f6c\u6362\u7ed3\u679c\u6587\u4ef6\u672a\u627e\u5230"}

    out_name = os.path.basename(out_path)
    out_size_kb = round(os.path.getsize(out_path) / 1024, 1)
    is_static = out_path.startswith(DOWNLOAD_DIR)
    if is_static:
        file_url = f"/static/downloads/{out_name}"
    else:
        file_url = f"/api/download/{out_name}"

    pages_info = (
        f"\uff0c\u5171 {result['pages']} \u9875"
        if "pages" in result else ""
    )
    tgt_label = {
        "pdf": "PDF",
        "docx": "Word",
        "jpeg": "JPG",
        "png": "PNG",
        "webp": "WebP",
    }.get(tgt_fmt, tgt_fmt.upper())
    src_name = os.path.basename(file_path)

    content = (
        "\u2705 **\u8f6c\u6362\u5b8c\u6210**\n\n"
        f"- \u6e90\u6587\u4ef6\uff1a{src_name}\n"
        f"- \u76ee\u6807\uff1a{tgt_label}{pages_info}\n"
        f"- \u5927\u5c0f\uff1a{out_size_kb} KB\n\n"
        f"[\u4e0b\u8f7d\u6587\u4ef6]({file_url})"
    )
    return {"file_url": file_url, "content": content}
